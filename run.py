#!/usr/bin/env python3
"""
Automated Job Matching + Resume Tailoring Bot
Runs daily, fetches remote jobs, scores them, generates tailored docs, sends email.
"""

import os
import sys
import json
import re
import zipfile
import smtplib
import logging
import traceback
from datetime import date
from email.message import EmailMessage
from pathlib import Path
from io import BytesIO

import requests
import yaml
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
import feedparser

# ── Logging ──────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.StreamHandler(sys.stdout)],
)
log = logging.getLogger(__name__)

# ── Paths ─────────────────────────────────────────────────────────────────────
ROOT = Path(__file__).parent
OUTPUT = ROOT / "output"
OUTPUT.mkdir(exist_ok=True)
TEMPLATES = ROOT / "templates"
BULLET_BANK_PATH = ROOT / "bullet_bank.json"
CONFIG_PATH = ROOT / "config.yaml"

MAX_JOB_TEXT = 20_000
TOP_N = 5
BULLET_COUNT = 10


# ══════════════════════════════════════════════════════════════════════════════
# CONFIG + DATA LOADING
# ══════════════════════════════════════════════════════════════════════════════

def load_config() -> dict:
    with open(CONFIG_PATH) as f:
        return yaml.safe_load(f)


def load_bullet_bank() -> dict:
    with open(BULLET_BANK_PATH) as f:
        return json.load(f)


# ══════════════════════════════════════════════════════════════════════════════
# JOB FETCHING
# ══════════════════════════════════════════════════════════════════════════════

HEADERS = {"User-Agent": "JobBot/1.0 (automated job matching; contact via GitHub)"}


def fetch_remotive(cfg: dict) -> list[dict]:
    jobs = []
    try:
        r = requests.get("https://remotive.com/api/remote-jobs", headers=HEADERS, timeout=20)
        r.raise_for_status()
        data = r.json()
        for j in data.get("jobs", []):
            jobs.append({
                "title": j.get("title", ""),
                "company": j.get("company_name", ""),
                "url": j.get("url", ""),
                "description": j.get("description", "")[:MAX_JOB_TEXT],
                "salary_raw": j.get("salary", ""),
                "source": "Remotive",
            })
        log.info(f"Remotive: fetched {len(jobs)} jobs")
    except Exception as e:
        log.warning(f"Remotive fetch failed: {e}")
    return jobs


def fetch_weworkremotely() -> list[dict]:
    jobs = []
    try:
        feed = feedparser.parse("https://weworkremotely.com/remote-jobs.rss")
        for entry in feed.entries:
            desc_html = entry.get("summary", "") or entry.get("content", [{}])[0].get("value", "")
            desc_text = BeautifulSoup(desc_html, "lxml").get_text(" ", strip=True)[:MAX_JOB_TEXT]
            jobs.append({
                "title": entry.get("title", ""),
                "company": entry.get("author", "") or _extract_company_from_title(entry.get("title", "")),
                "url": entry.get("link", ""),
                "description": desc_text,
                "salary_raw": "",
                "source": "WeWorkRemotely",
            })
        log.info(f"WeWorkRemotely: fetched {len(jobs)} jobs")
    except Exception as e:
        log.warning(f"WeWorkRemotely fetch failed: {e}")
    return jobs


def fetch_remoteok() -> list[dict]:
    jobs = []
    try:
        r = requests.get("https://remoteok.io/api", headers=HEADERS, timeout=20)
        r.raise_for_status()
        data = r.json()
        for j in data[1:] if isinstance(data, list) else []:
            if not isinstance(j, dict):
                continue
            desc_html = j.get("description", "")
            desc_text = BeautifulSoup(desc_html, "lxml").get_text(" ", strip=True)[:MAX_JOB_TEXT] if desc_html else ""
            jobs.append({
                "title": j.get("position", ""),
                "company": j.get("company", ""),
                "url": j.get("url", "") or f"https://remoteok.io/remote-jobs/{j.get('id', '')}",
                "description": desc_text,
                "salary_raw": f"{j.get('salary_min', '')} - {j.get('salary_max', '')}",
                "source": "RemoteOK",
            })
        log.info(f"RemoteOK: fetched {len(jobs)} jobs")
    except Exception as e:
        log.warning(f"RemoteOK fetch failed: {e}")
    return jobs


def _extract_company_from_title(title: str) -> str:
    if ":" in title:
        return title.split(":")[0].strip()
    return ""


def dedupe_jobs(jobs: list[dict]) -> list[dict]:
    seen = set()
    out = []
    for j in jobs:
        key = j["url"].strip().rstrip("/")
        if key and key not in seen:
            seen.add(key)
            out.append(j)
    return out


# ══════════════════════════════════════════════════════════════════════════════
# FILTERING
# ══════════════════════════════════════════════════════════════════════════════

def extract_salary(text: str, salary_raw: str) -> int | None:
    combined = f"{salary_raw} {text[:3000]}".replace(",", "")
    patterns = [
        r"\$\s*(\d{3,6})k",
        r"\$\s*(\d{6})",
        r"(\d{3})[\s,]*000\s*(?:USD|usd|/yr|/year|annually)?",
        r"USD\s*(\d{3,6})k",
    ]
    values = []
    for pat in patterns:
        for m in re.finditer(pat, combined, re.IGNORECASE):
            try:
                val = int(m.group(1))
                if "k" in pat.lower():
                    val *= 1000
                if 50_000 <= val <= 1_000_000:
                    values.append(val)
            except Exception:
                pass
    return max(values) if values else None


def has_bonus_or_equity(text: str) -> bool:
    signals = ["bonus", "equity", "stock", "rsu", "options", "401k", "401(k)", "profit sharing"]
    lower = text.lower()
    return any(s in lower for s in signals)


def is_acceptable_location(text: str, title: str) -> bool:
    """
    Accept: fully remote, remote-first, hybrid (penalized in score but not rejected),
    and on-site at specific preferred companies.
    Reject: hard on-site only with relocation required.
    """
    lower = (text + " " + title).lower()
    # Hard reject — explicitly requires relocation or is clearly on-site only
    hard_reject = ["relocation required", "must relocate", "in-office only", "on-site only"]
    for r in hard_reject:
        if r in lower:
            return False
    # All sources are remote-focused, so accept by default
    return True


def is_fulltime(text: str) -> bool:
    lower = text.lower()
    if any(x in lower for x in ["contract", "freelance", "part-time", "part time", "temporary"]):
        return False
    return True


def passes_filters(job: dict, cfg: dict) -> tuple[bool, str]:
    title = job["title"]
    desc = job["description"]
    salary_raw = job["salary_raw"]
    combined_text = f"{title} {desc}"

    if not is_fulltime(combined_text):
        return False, "not full-time"

    if not is_acceptable_location(desc, title):
        return False, "relocation required / hard on-site"

    salary = extract_salary(desc, salary_raw)
    min_salary = cfg.get("filters", {}).get("min_salary", 150_000)
    if salary is not None and salary < min_salary:
        return False, f"salary {salary} < {min_salary}"

    if not has_bonus_or_equity(combined_text):
        return False, "no bonus/equity mention"

    return True, ""


# ══════════════════════════════════════════════════════════════════════════════
# SCORING
# ══════════════════════════════════════════════════════════════════════════════

def score_job(job: dict, bullet_bank: dict, cfg: dict) -> float:
    title = job["title"].lower()
    desc = job["description"].lower()
    combined = f"{title} {desc}"

    score = 0.0

    # 0. Title blocklist
    for blocked in cfg.get("scoring", {}).get("title_blocklist", []):
        if blocked.lower() in title:
            score -= 20.0
            break

    # 1. Keywords priority overlap (+2 each)
    kw_priority = bullet_bank.get("keywords_priority", [])
    kw_hits = sum(1 for kw in kw_priority if kw.lower() in combined)
    score += kw_hits * 2.0

    # 2. Title signals (+5 each)
    title_signals = cfg.get("scoring", {}).get("title_signals", [])
    title_hits = sum(1 for t in title_signals if t.lower() in title)
    score += title_hits * 5.0

    # 3. Industry signals (+3 each, stacking)
    industry_signals = cfg.get("scoring", {}).get("industry_signals", [])
    industry_hits_list = [s for s in industry_signals if s.lower() in combined]
    score += len(industry_hits_list) * 3.0
    if industry_hits_list:
        job["_industry_hits"] = industry_hits_list[:6]

    # 4. Remote preference — boost fully remote, don't penalize hybrid
    lower_combined = combined
    if any(s in lower_combined for s in ["fully remote", "100% remote", "remote-first", "work from anywhere", "remote only"]):
        score += 5.0
    elif any(s in lower_combined for s in ["hybrid", "in-office"]):
        score -= 3.0  # slight penalty but don't reject

    # 5. Comp signals
    salary = extract_salary(desc, job["salary_raw"])
    min_salary = cfg.get("filters", {}).get("min_salary", 150_000)
    if salary and salary >= min_salary:
        score += 5.0
    if has_bonus_or_equity(combined):
        score += 2.0

    # 6. Bullet bank keyword overlap (+0.5 each)
    bullet_overlap = 0
    for b in bullet_bank.get("bullets", []):
        hits = sum(1 for k in b.get("keywords", []) if k.lower() in combined)
        bullet_overlap += hits
    score += bullet_overlap * 0.5

    return round(score, 2)


def select_bullets(job: dict, bullet_bank: dict, n: int = BULLET_COUNT) -> list[str]:
    desc = job["description"].lower()
    scored = []
    for b in bullet_bank.get("bullets", []):
        hits = sum(1 for k in b.get("keywords", []) if k.lower() in desc)
        scored.append((hits, b.get("text", "")))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [text for _, text in scored[:n]]


# ══════════════════════════════════════════════════════════════════════════════
# DOCX GENERATION
# ══════════════════════════════════════════════════════════════════════════════

def load_master_resume_info() -> dict:
    defaults = {
        "name": "YOUR NAME",
        "email": "your@email.com",
        "phone": "555-555-5555",
        "location": "Remote, USA",
        "linkedin": "linkedin.com/in/yourprofile",
    }
    master = ROOT / "master_resume.docx"
    if not master.exists():
        log.warning("master_resume.docx not found — using placeholder defaults")
        return defaults
    try:
        doc = Document(master)
        for para in doc.paragraphs:
            if para.text.strip():
                defaults["name"] = para.text.strip()
                break
    except Exception as e:
        log.warning(f"Could not read master_resume.docx: {e}")
    return defaults


def fill_template(template_path: Path, replacements: dict) -> Document:
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")
    doc = Document(template_path)

    def replace_in_runs(paragraph, replacements):
        for run in paragraph.runs:
            for key, value in replacements.items():
                token = "{{" + key + "}}"
                if token in run.text:
                    run.text = run.text.replace(token, str(value))
        full = paragraph.text
        for key, value in replacements.items():
            token = "{{" + key + "}}"
            if token in full:
                for run in paragraph.runs:
                    run.text = ""
                if paragraph.runs:
                    paragraph.runs[0].text = full.replace(token, str(value))
                break

    for para in doc.paragraphs:
        replace_in_runs(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_runs(para, replacements)
    return doc


def build_cover_body(job: dict, matched_keywords: list[str], person: dict) -> str:
    kw_str = ", ".join(matched_keywords[:6]) if matched_keywords else "your key requirements"
    industry_hits = job.get("_industry_hits", [])

    # Personalized opening based on detected vertical
    industry_line = ""
    ih_lower = [h.lower() for h in industry_hits]
    if any(k in ih_lower for k in ["artificial intelligence", "machine learning", "llm", "generative ai", "mlops", "ai product", "data platform", "data streaming", "kafka"]):
        industry_line = (
            f"I am particularly drawn to {job['company']}'s work at the intersection of AI, data infrastructure, "
            f"and enterprise technology — and to the opportunity to drive program and product delivery in this space."
        )
    elif any(k in ih_lower for k in ["whoop", "wearable", "connected fitness", "fitness platform", "human performance", "biometric", "sports technology", "athlete performance"]):
        industry_line = (
            f"I am especially excited by {job['company']}'s mission in human performance and connected fitness — "
            f"the intersection of technology, data, and athletic performance is exactly where I want to build."
        )
    elif any(k in ih_lower for k in ["sports betting", "fantasy sports", "draftkings", "fanduel", "betmgm", "sports analytics"]):
        industry_line = (
            f"Your work at the crossroads of sports and technology is exactly the kind of high-energy, data-driven "
            f"environment I thrive in — and I am excited about the opportunity to contribute at {job['company']}."
        )
    elif any(k in ih_lower for k in ["digital health", "healthtech", "wellness platform"]):
        industry_line = (
            f"I am drawn to {job['company']}'s mission in digital health and the meaningful impact your platform "
            f"has on how people manage their health and wellbeing."
        )
    elif any(k in ih_lower for k in ["saas", "b2b software", "fintech", "enterprise software", "api platform"]):
        industry_line = (
            f"Your position as a leader in the enterprise SaaS and technology space is a strong draw — "
            f"I have spent my career driving program and product outcomes in exactly this environment."
        )

    body = (
        f"I am writing to express my strong interest in the {job['title']} position at {job['company']}. "
        f"With a proven track record in technical program and product management, I am confident in my ability "
        f"to drive meaningful results for your team.\n\n"
        f"{industry_line + chr(10) + chr(10) if industry_line else ''}"
        f"My experience aligns directly with your needs in areas such as {kw_str}. "
        f"I have consistently led cross-functional teams, navigated ambiguity, "
        f"and delivered high-impact programs on time and within budget — at organizations ranging from "
        f"high-growth startups to Fortune 500 enterprises.\n\n"
        f"I bring a strong bias for action, structured thinking, and the ability to bridge deeply technical "
        f"teams with business stakeholders at the executive level. I would welcome the opportunity to discuss "
        f"how my background can contribute to {job['company']}'s continued growth.\n\n"
        f"Thank you for your consideration. I look forward to speaking with you."
    )
    return body


def generate_resume(job: dict, bullets: list[str], person: dict, idx: int) -> Path:
    template = TEMPLATES / "resume_template.docx"
    replacements = {
        "NAME": person["name"],
        "EMAIL": person.get("email", ""),
        "PHONE": person.get("phone", ""),
        "LOCATION": person.get("location", "Remote"),
        "LINKEDIN": person.get("linkedin", ""),
        "TITLE": job["title"],
        "COMPANY": job["company"],
        "DATE": date.today().strftime("%B %Y"),
    }
    for i, bullet in enumerate(bullets, 1):
        replacements[f"BULLET_{i}"] = bullet
    for i in range(len(bullets) + 1, BULLET_COUNT + 1):
        replacements[f"BULLET_{i}"] = ""

    doc = fill_template(template, replacements)
    safe_company = re.sub(r"[^\w]", "_", job["company"])[:30]
    filename = OUTPUT / f"resume_{idx:02d}_{safe_company}.docx"
    doc.save(filename)
    return filename


def generate_cover_letter(job: dict, bullets: list[str], person: dict, idx: int) -> Path:
    template = TEMPLATES / "cover_letter_template.docx"
    bullet_bank = load_bullet_bank()
    kw_priority = bullet_bank.get("keywords_priority", [])
    desc_lower = job["description"].lower()
    matched_kws = [k for k in kw_priority if k.lower() in desc_lower][:8]

    cover_body = build_cover_body(job, matched_kws, person)
    replacements = {
        "NAME": person["name"],
        "DATE": date.today().strftime("%B %d, %Y"),
        "COMPANY": job["company"],
        "ROLE": job["title"],
        "COVER_BODY": cover_body,
        "EMAIL": person.get("email", ""),
        "PHONE": person.get("phone", ""),
        "LINKEDIN": person.get("linkedin", ""),
    }
    doc = fill_template(template, replacements)
    safe_company = re.sub(r"[^\w]", "_", job["company"])[:30]
    filename = OUTPUT / f"cover_letter_{idx:02d}_{safe_company}.docx"
    doc.save(filename)
    return filename


# ══════════════════════════════════════════════════════════════════════════════
# SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

def write_summary(ranked: list[dict]) -> Path:
    lines = [
        f"Job Match Summary — {date.today().strftime('%A, %B %d, %Y')}",
        "=" * 60,
        "",
    ]
    for i, job in enumerate(ranked, 1):
        salary = extract_salary(job["description"], job["salary_raw"])
        sal_str = f"${salary:,}" if salary else "Not stated"
        industry_hits = job.get("_industry_hits", [])
        lines += [
            f"#{i}  {job['title']} @ {job['company']}",
            f"    Score: {job['_score']}  |  Salary: {sal_str}  |  Source: {job['source']}",
            f"    URL: {job['url']}",
            f"    Why: {job.get('_why', 'Strong keyword overlap')}",
        ]
        if industry_hits:
            lines.append(f"    Industry signals: {', '.join(industry_hits)}")
        lines.append("")
    path = OUTPUT / "summary.txt"
    path.write_text("\n".join(lines))
    return path


def build_why(job: dict, bullet_bank: dict, cfg: dict) -> str:
    desc = job["description"].lower()
    kw_priority = bullet_bank.get("keywords_priority", [])
    hits = [k for k in kw_priority if k.lower() in desc][:5]
    salary = extract_salary(desc, job["salary_raw"])
    industry_hits = job.get("_industry_hits", [])
    parts = []
    if hits:
        parts.append(f"keywords: {', '.join(hits)}")
    if salary:
        parts.append(f"salary ${salary:,}")
    if has_bonus_or_equity(desc):
        parts.append("bonus/equity")
    if industry_hits:
        parts.append(f"industry: {', '.join(industry_hits[:3])}")
    return "; ".join(parts) if parts else "general remote match"


# ══════════════════════════════════════════════════════════════════════════════
# ZIP
# ══════════════════════════════════════════════════════════════════════════════

def create_zip(files: list[Path], summary: Path) -> Path:
    zip_path = OUTPUT / "tailored_packets.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.write(summary, summary.name)
        for f in files:
            if f.exists():
                zf.write(f, f.name)
    log.info(f"ZIP created: {zip_path}")
    return zip_path


# ══════════════════════════════════════════════════════════════════════════════
# EMAIL
# ══════════════════════════════════════════════════════════════════════════════

def send_email(summary_path: Path, attachments: list[Path], zip_path: Path, ranked: list[dict]):
    smtp_user = os.environ.get("GMAIL_SMTP_USER", "")
    smtp_pass = os.environ.get("GMAIL_APP_PASSWORD", "")
    email_to = os.environ.get("EMAIL_TO", smtp_user)

    if not smtp_user or not smtp_pass:
        log.warning("Email credentials not set. Skipping email send.")
        return False

    body_lines = [
        f"Good morning! Here are your Top {TOP_N} job matches for {date.today().strftime('%B %d, %Y')}.",
        "",
    ]
    for i, job in enumerate(ranked, 1):
        salary = extract_salary(job["description"], job["salary_raw"])
        sal_str = f"${salary:,}" if salary else "Not stated"
        industry_hits = job.get("_industry_hits", [])
        ind_str = f" | Industry: {', '.join(industry_hits[:3])}" if industry_hits else ""
        body_lines += [
            f"#{i}. {job['title']} @ {job['company']}",
            f"   Score: {job['_score']} | Salary: {sal_str} | Source: {job['source']}{ind_str}",
            f"   {job['url']}",
            f"   Why: {job.get('_why', '')}",
            "",
        ]
    body_lines += ["Attachments: tailored resumes + cover letters + ZIP.", "", "Good luck! 🚀"]
    body = "\n".join(body_lines)

    msg = EmailMessage()
    msg["Subject"] = f"🎯 Top {TOP_N} Job Matches — {date.today().strftime('%b %d, %Y')}"
    msg["From"] = smtp_user
    msg["To"] = email_to
    msg.set_content(body)

    for fpath in list(attachments) + [zip_path, summary_path]:
        if fpath and fpath.exists():
            with open(fpath, "rb") as f:
                data = f.read()
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if fpath.suffix == ".zip":
                mime = "application/zip"
            elif fpath.suffix == ".txt":
                mime = "text/plain"
            msg.add_attachment(data, maintype=mime.split("/")[0],
                               subtype=mime.split("/")[1], filename=fpath.name)

    sent = False
    for port, use_ssl in [(587, False), (465, True)]:
        try:
            if use_ssl:
                server = smtplib.SMTP_SSL("smtp.gmail.com", port, timeout=30)
            else:
                server = smtplib.SMTP("smtp.gmail.com", port, timeout=30)
                server.ehlo()
                server.starttls()
                server.ehlo()
            server.login(smtp_user, smtp_pass)
            server.send_message(msg)
            server.quit()
            log.info(f"Email sent via port {port} to {email_to}")
            sent = True
            break
        except Exception as e:
            log.warning(f"Email attempt (port {port}) failed: {e}")

    if not sent:
        log.error("All email send attempts failed. Artifacts will still be uploaded.")
    return sent


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════

def main():
    log.info("=== Job Bot Starting ===")
    cfg = load_config()
    bullet_bank = load_bullet_bank()
    person = load_master_resume_info()

    log.info("Fetching jobs...")
    jobs = []
    jobs += fetch_remotive(cfg)
    jobs += fetch_weworkremotely()
    jobs += fetch_remoteok()
    log.info(f"Total fetched: {len(jobs)} jobs")

    jobs = dedupe_jobs(jobs)
    log.info(f"After dedup: {len(jobs)} jobs")

    passing = []
    for job in jobs:
        ok, reason = passes_filters(job, cfg)
        if ok:
            passing.append(job)
    log.info(f"After filters: {len(passing)} jobs pass")

    if not passing:
        log.warning("No jobs passed filters. Consider loosening config.yaml filters.")
        summary_path = OUTPUT / "summary.txt"
        summary_path.write_text(
            f"No jobs matched filters on {date.today()}.\n"
            "Consider loosening min_salary or filter criteria in config.yaml."
        )
        create_zip([], summary_path)
        return

    for job in passing:
        job["_score"] = score_job(job, bullet_bank, cfg)
        job["_why"] = build_why(job, bullet_bank, cfg)

    ranked = sorted(passing, key=lambda j: j["_score"], reverse=True)[:TOP_N]
    log.info(f"Top {len(ranked)} matches selected")
    for i, j in enumerate(ranked, 1):
        ind = j.get("_industry_hits", [])
        log.info(f"  #{i} {j['title']} @ {j['company']} — score {j['_score']}" +
                 (f" | industry: {', '.join(ind[:3])}" if ind else ""))

    resume_files = []
    cover_files = []
    for idx, job in enumerate(ranked, 1):
        bullets = select_bullets(job, bullet_bank)
        try:
            rf = generate_resume(job, bullets, person, idx)
            resume_files.append(rf)
            log.info(f"  Resume generated: {rf.name}")
        except Exception as e:
            log.error(f"Resume generation failed for job {idx}: {e}\n{traceback.format_exc()}")
        try:
            cf = generate_cover_letter(job, bullets, person, idx)
            cover_files.append(cf)
            log.info(f"  Cover letter generated: {cf.name}")
        except Exception as e:
            log.error(f"Cover letter generation failed for job {idx}: {e}\n{traceback.format_exc()}")

    summary_path = write_summary(ranked)
    log.info(f"Summary written: {summary_path}")

    zip_path = create_zip(resume_files + cover_files, summary_path)

    try:
        send_email(summary_path, resume_files + cover_files, zip_path, ranked)
    except Exception as e:
        log.error(f"Email send raised unexpected exception: {e}")

    log.info("=== Job Bot Complete ===")


if __name__ == "__main__":
    main()
